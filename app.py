import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import google.generativeai as genai
import os
import re
from collections import Counter
import io

# ==========================================
# 🎨 1. 브레인큐브 브랜드 테마 설정 및 로고
# ==========================================
st.set_page_config(page_title="Brain Cube AI Universal Report", page_icon="braincube_logo.png", layout="wide", initial_sidebar_state="expanded")

logo_filename = "braincube_logo.png"

if os.path.exists(logo_filename):
    with st.sidebar:
        st.image(logo_filename, use_container_width=True)
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #777777; font-size: 0.9rem;'>the best company true partner</p>", unsafe_allow_html=True)
        st.divider()

BRAIN_CUBE_ORANGE = "#F7931D"
BRAIN_CUBE_TEXT = "#555555"
CHART_COLORS = [BRAIN_CUBE_ORANGE, '#1F77B4', '#2CA02C', 
'#D62728', '#9467BD', '#8C564B', '#E377C2', '#7F7F7F', '#BCBD22', '#17BECF']

st.markdown(f"""
<style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
html, body, div, p, span, h1, h2, h3, h4, h5, h6, label, button, input {{
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, system-ui, Roboto, 'Helvetica Neue', sans-serif;
}}
    [class*="material-symbols"], [class*="icon"], .material-icons {{
        font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
font-feature-settings: 'liga' 1;
    }}
    .stApp {{ background-color: #FAFAFA;
}}
    .stMarkdown h1 {{ background: linear-gradient(135deg, {BRAIN_CUBE_ORANGE} 0%, #FFB75E 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 900;
letter-spacing: -1px; }}
    .stMarkdown h2 {{ color: {BRAIN_CUBE_ORANGE}; margin-top: 1.5rem; font-weight: 800;
}}
    .stDivider hr {{ border-color: {BRAIN_CUBE_ORANGE}; opacity: 0.2; }}
    [data-testid="metric-container"] {{ background-color: #FFFFFF;
border: 1px solid #EBEBEB; padding: 15px 10px; border-radius: 12px; box-shadow: 0 4px 10px rgba(0, 0, 0, 0.04);
border-left: 6px solid {BRAIN_CUBE_ORANGE}; transition: transform 0.2s ease-in-out; }}
    [data-testid="metric-container"]:hover {{ transform: translateY(-3px);
box-shadow: 0 6px 15px rgba(247, 147, 29, 0.15); }}
    .stMetric [data-testid="stMetricLabel"] {{ color: #777777; font-weight: 600;
font-size: 1.0rem; }}
    .stMetric [data-testid="stMetricValue"] {{ color: #333333 !important; font-weight: 800 !important; font-size: clamp(1.1rem, 1.5vw, 1.7rem) !important;
white-space: nowrap !important; word-break: keep-all !important; letter-spacing: -0.5px; }}
    .stButton button {{ background: linear-gradient(135deg, {BRAIN_CUBE_ORANGE} 0%, #E07D10 100%);
color: white; border: none; border-radius: 8px; padding: 12px 24px; font-weight: 700; box-shadow: 0 4px 6px rgba(247, 147, 29, 0.3);
transition: all 0.3s; }}
    .stButton button:hover {{ transform: translateY(-2px); box-shadow: 0 6px 12px rgba(247, 147, 29, 0.4);
}}
    .streamlit-expanderHeader {{ background-color: #FFF8F0; border-radius: 8px; font-weight: 600; color: #444444;
}}
    .stSidebar [data-testid="stMarkdownContainer"] p {{ color: {BRAIN_CUBE_TEXT}; }}
    [data-testid="stDataFrame"] {{ width: 100% !important;
}}
</style>
""", unsafe_allow_html=True)

# ==========================================
# ⚙️ 2. AI 모델 및 벤치마크 데이터 설정
# ==========================================

api_key = os.environ.get("GOOGLE_API_KEY", "")
if not api_key:
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
    except (KeyError, FileNotFoundError):
        pass
if not api_key:
    # fallback — 배포 시 반드시 환경변수 또는 secrets.toml 사용
    api_key = "AIzaSyCVbSKxEoE3KBag7G6BPFmPFImkn4rkykQ"
genai.configure(api_key=api_key)

vision_model_name = None
for m in genai.list_models():
    if 'generateContent' in m.supported_generation_methods:
        if '1.5' in m.name or 'vision' in m.name:
            vision_model_name = m.name; break
if vision_model_name is None:
    for m in genai.list_models():
        if 'generateContent' in m.supported_generation_methods:
            vision_model_name = m.name; break
model = genai.GenerativeModel(vision_model_name)

benchmark_df = None
industry_list = ["전체 업종 (평균)"]
try:
    from modules.firebase_connector import load_benchmark as _fb_load
    benchmark_df = _fb_load()
    if benchmark_df is not None and '분야' in benchmark_df.columns:
        industry_list += sorted(benchmark_df['분야'].dropna().astype(str).unique().tolist())
except Exception as e: pass

if "ai_report_content" not in st.session_state: st.session_state.ai_report_content = ""
if "saved_charts" not in st.session_state: st.session_state.saved_charts = {}

# ==========================================
# 📊 3. 메인 화면 UI 및 파일 로드
# ==========================================
st.title("Brain Cube Universal AI Report")
st.markdown("현재 캠페인의 성과를 다각도로 시각화하고 26년 브레인큐브 평균과 정밀 분석합니다.")
st.divider()

uploaded_files = st.file_uploader("캠페인 엑셀 파일(.xlsx, .csv)과 소재 이미지(선택) 업로드", type=['xlsx', 'csv', 'png', 'jpg', 'jpeg'], accept_multiple_files=True)
excel_file, image_files = None, {}

if uploaded_files:
    for file in uploaded_files:
        if (file.name.endswith('.xlsx') or file.name.endswith('.csv')) and excel_file is None: excel_file = file
        elif not file.name.endswith('.xlsx') and not file.name.endswith('.csv'): image_files[os.path.splitext(file.name)[0]] = file

if excel_file is not None:
    file_name = excel_file.name 
    try:
        if file_name.endswith('.xlsx'):
            xl = pd.ExcelFile(excel_file)
            target_sheet = xl.sheet_names[0]
            for sheet in xl.sheet_names:
                if sheet.strip().upper() in ['RAW', 'RAW DATA']: target_sheet = sheet; break
            df = pd.read_excel(excel_file, sheet_name=target_sheet)
        else: df = pd.read_csv(excel_file)
    except Exception as e: st.error(f"오류 발생: {e}"); st.stop()
    
    with st.expander("⚙️ 데이터 항목 매칭 설정 (클릭하여 수정)"):
        cols = df.columns.tolist()
        cols_with_none = ["(항목 없음)"] + cols
        def get_idx(kw_list):
            c_clean = [str(c).lower().replace(" ", "") for c in cols]
            for kw in kw_list:
                for i, c in enumerate(c_clean):
                    if kw == c: return i
            for kw in kw_list:
                for i, c in enumerate(c_clean):
                    if kw in c and all(x not in c for x in ['률', '율', '%', 'rate']): return i
            for kw in kw_list:
                for i, c in enumerate(c_clean):
                    if kw in c: return i
            return 0
        def get_idx_opt(kw_list):
            c_clean = [str(c).lower().replace(" ", "") for c in cols]
            for kw in kw_list:
                for i, c in enumerate(c_clean):
                    if kw == c: return i + 1
            for kw in kw_list:
                for i, c in enumerate(c_clean):
                    if kw in c and all(x not in c for x in ['률', '율', '%', 'rate']): return i + 1
            for kw in kw_list:
                for i, c in enumerate(c_clean):
                    if kw in c: return i + 1
            return 0

        c1, c2, c3, c4 = st.columns(4)
        col_date = c1.selectbox("📅 날짜", cols, index=get_idx(['date', '일자', '날짜']))
        col_media = c2.selectbox("📢 매체(채널)", cols, index=get_idx(['매체', '채널']))
        col_creative = c3.selectbox("📝 소재(발송제목)", cols_with_none, index=get_idx_opt(['소재명', '소재', '제목', '발송제목', '문구', '광고명', '캠페인명']))
        col_cost = c4.selectbox("💰 광고비(집행금액)", cols, index=get_idx(['cost_raw', '광고비', '집행금액', '소진액', '사용금액', '비용', 'cost', '금액']))

        c5, c6, c7, c8 = st.columns(4)
        col_imp = c5.selectbox("👁️ 노출/발송수", cols_with_none, index=get_idx_opt(['노출수', '노출량', '발송수', '발송건', '발송수량', '발송', 'imp', '도달', '뷰', 'view', 'reach', '조회', '노출']))
        col_clicks = c6.selectbox("👆 클릭수", cols_with_none, index=get_idx_opt(['클릭수', '클릭건', 'click', '클릭', '방문', '유입']))
        col_db = c7.selectbox("🎯 DB/전환수", cols_with_none, index=get_idx_opt(['주문수', '주문건수', '주문', '구매수', '구매', '결제', '전환수', '전환건', '전환', 'db', '가입', '신청', '완료', '실적']))
    
    st.divider()

    try:
        df[col_date] = pd.to_datetime(df[col_date], errors='coerce')
        df = df.dropna(subset=[col_date])

        st.subheader("🔍 분석 세부 설정")
        col_ui_1, col_ui_2, col_ui_3 = st.columns(3)
        with col_ui_1: selected_date_range = st.date_input("📅 조회 기간", value=(df[col_date].min().date(), df[col_date].max().date()))
        with col_ui_2: selected_media = st.multiselect("🔎 매체 필터링", options=df[col_media].astype(str).unique().tolist(), default=[])
        with col_ui_3: selected_industry = st.selectbox("🏢 26년 비교 업종", options=industry_list)
        
        # 💡 [핵심 업데이트] CTR 목표 입력칸 추가 (3열 레이아웃으로 변경)
        with st.expander("🎯 캠페인 목표 KPI 설정 (계기판 출력용)"):
            kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
            target_budget = kpi_col1.number_input("이번 캠페인 총 목표 예산 (원)", min_value=0, value=0, step=100000)
            target_cpa = kpi_col2.number_input("목표 Target CPA (원)", min_value=0, value=0, step=1000)
            target_ctr = kpi_col3.number_input("목표 Target CTR (%)", min_value=0.0, value=0.0, step=0.1, format="%.2f")

        with st.expander("📂 업로드된 원본 데이터 전체 보기 (스크롤 지원)"):
            st.dataframe(df, use_container_width=True, height=250)
            
        st.divider()

        if len(selected_date_range) == 2:
            start_date, end_date = selected_date_range
            mask = (df[col_date].dt.date >= start_date) & (df[col_date].dt.date <= end_date)
            filtered_df = df.loc[mask].copy()

            if selected_media: filtered_df = filtered_df[filtered_df[col_media].isin(selected_media)]

            for col in [col_cost] + ([col_imp] if col_imp != "(항목 없음)" else []) + ([col_clicks] if col_clicks != "(항목 없음)" else []) + ([col_db] if col_db != "(항목 없음)" else []):
                filtered_df[col] = pd.to_numeric(filtered_df[col].astype(str).str.replace(',', '').str.replace('-', ''), errors='coerce').fillna(0)

            total_spent = filtered_df[col_cost].sum()
            total_imp = filtered_df[col_imp].sum() if col_imp != "(항목 없음)" else 0
            total_clicks = filtered_df[col_clicks].sum() if col_clicks != "(항목 없음)" else 0
            total_db = filtered_df[col_db].sum() if col_db != "(항목 없음)" else 0
            
            has_db = col_db != "(항목 없음)" and total_db > 0
            has_clicks = col_clicks != "(항목 없음)" and total_clicks > 0
            has_imp = col_imp != "(항목 없음)" and total_imp > 0

            current_cpa = total_spent / total_db if has_db and total_db > 0 else 0
            current_ctr = (total_clicks / total_imp) * 100 if has_clicks and has_imp and total_imp > 0 else 0

            if has_db: main_metric, is_ascending, efficiency_value = "CPA", True, f"{current_cpa:,.0f} 원"
            elif has_clicks and has_imp: main_metric, is_ascending, efficiency_value = "CTR", False, f"{current_ctr:.2f} %"
            elif has_clicks: main_metric, is_ascending, efficiency_value = "CPC", True, f"{total_spent/total_clicks:,.0f} 원"
            else: main_metric, is_ascending, efficiency_value = col_cost, False, "데이터 부족"
                
            bench_ctr_avg, bench_cpa_avg, bench_metric_delta = 0, 0, None
            temp_bench = None
            if benchmark_df is not None:
                temp_bench = benchmark_df.copy()
                if selected_industry != "전체 업종 (평균)" and '분야' in temp_bench.columns: temp_bench = temp_bench[temp_bench['분야'] == selected_industry]
                if selected_media and '매체' in temp_bench.columns: temp_bench = temp_bench[temp_bench['매체'].isin(selected_media)]
                
            if '발송건' in temp_bench.columns and '클릭수' in temp_bench.columns and temp_bench['발송건'].sum() > 0: bench_ctr_avg = (temp_bench['클릭수'].sum() / temp_bench['발송건'].sum()) * 100
            if '광고비' in temp_bench.columns and 'DB' in temp_bench.columns and temp_bench['DB'].sum() > 0: bench_cpa_avg = temp_bench['광고비'].sum() / temp_bench['DB'].sum()
                
            if main_metric == "CPA" and bench_cpa_avg > 0: bench_metric_delta = f"{current_cpa - bench_cpa_avg:,.0f} 원 (26년 평균 대비)"
            elif main_metric == "CTR" and bench_ctr_avg > 0: bench_metric_delta = f"{current_ctr - bench_ctr_avg:.2f}% (26년 평균 대비)"

            st.session_state.saved_charts = {}
            st.subheader(f"💡 브레인큐브 핵심 퍼포먼스 지표")
            
            # ==========================================
            # 🚨 [업데이트 완료] 활성화된 목표 개수에 따라 계기판을 동적으로 배치합니다!
            # ==========================================
            active_gauges = []
            if target_budget > 0: active_gauges.append('budget')
            if has_db and target_cpa > 0: active_gauges.append('cpa')
            if has_clicks and has_imp and target_ctr > 0: active_gauges.append('ctr')

            if active_gauges:
                g_cols = st.columns(len(active_gauges))
                col_idx = 0
                
                # 1. 예산 집행 계기판
                if 'budget' in active_gauges:
                    fig_budget = go.Figure()
                    fig_budget.add_trace(go.Indicator(
                        mode="gauge+number+delta", 
                        value=total_spent, 
                        number={'valueformat': ',.0f'},
                        delta={'reference': target_budget, 'position': "top", 'valueformat': ',.0f'},
                        domain={'x': [0, 1], 'y': [0, 0.75]},
                        title={'text': "총 예산 집행 현황 (원)<br><span style='font-size:12px;color:gray'>마우스를 올려보세요</span>", 'font': {'size': 16}},
                        gauge={'axis': {'range': [None, target_budget]}, 'bar': {'color': BRAIN_CUBE_ORANGE},
                               'threshold': {'line': {'color': "red", 'width': 4}, 'thickness': 0.75, 'value': target_budget}}
                    ))
                    fig_budget.add_trace(go.Pie(
                        values=[1], textinfo='none', hoverinfo='text',
                        hovertext=f"💰 현재 집행금액: {total_spent:,.0f} 원<br>🎯 목표 예산: {target_budget:,.0f} 원<br>📊 집행율: {(total_spent/target_budget)*100 if target_budget>0 else 0:.1f}%",
                        hole=0.5, marker=dict(colors=['rgba(0,0,0,0)']), 
                        domain={'x': [0, 1], 'y': [0, 0.75]}, showlegend=False
                    ))
                    fig_budget.update_layout(height=300, margin=dict(l=20, r=20, t=50, b=20), hoverlabel=dict(bgcolor="white", font_size=14))
                    g_cols[col_idx].plotly_chart(fig_budget, use_container_width=True)
                    col_idx += 1
                
                # 2. CPA 계기판 (낮을수록 초록색)
                if 'cpa' in active_gauges:
                    fig_cpa = go.Figure()
                    fig_cpa.add_trace(go.Indicator(
                        mode="gauge+number+delta", 
                        value=current_cpa, 
                        number={'valueformat': ',.0f'},
                        delta={'reference': target_cpa, 'increasing': {'color': "red"}, 'decreasing': {'color': "green"}, 'valueformat': ',.0f'},
                        domain={'x': [0, 1], 'y': [0, 0.75]},
                        title={'text': "현재 CPA vs 목표 Target CPA<br><span style='font-size:12px;color:gray'>마우스를 올려보세요</span>", 'font': {'size': 16}},
                        gauge={'axis': {'range': [None, target_cpa * 1.5]}, 'bar': {'color': "green" if current_cpa <= target_cpa else "red"},
                               'threshold': {'line': {'color': BRAIN_CUBE_TEXT, 'width': 4}, 'thickness': 0.75, 'value': target_cpa}}
                    ))
                    fig_cpa.add_trace(go.Pie(
                        values=[1], textinfo='none', hoverinfo='text',
                        hovertext=f"🎯 현재 CPA: {current_cpa:,.0f} 원<br>🚩 목표 CPA: {target_cpa:,.0f} 원<br>💡 목표 대비: {current_cpa - target_cpa:,.0f} 원 ({'초과 🚨' if current_cpa > target_cpa else '달성 쾌조 ✅'})",
                        hole=0.5, marker=dict(colors=['rgba(0,0,0,0)']), 
                        domain={'x': [0, 1], 'y': [0, 0.75]}, showlegend=False
                    ))
                    fig_cpa.update_layout(height=300, margin=dict(l=20, r=20, t=50, b=20), hoverlabel=dict(bgcolor="white", font_size=14))
                    g_cols[col_idx].plotly_chart(fig_cpa, use_container_width=True)
                    col_idx += 1

                # 3. CTR 계기판 (높을수록 초록색)
                if 'ctr' in active_gauges:
                    fig_ctr = go.Figure()
                    fig_ctr.add_trace(go.Indicator(
                        mode="gauge+number+delta", 
                        value=current_ctr, 
                        number={'valueformat': '.2f'}, 
                        delta={'reference': target_ctr, 'increasing': {'color': "green"}, 'decreasing': {'color': "red"}, 'valueformat': '.2f'},
                        domain={'x': [0, 1], 'y': [0, 0.75]},
                        title={'text': "현재 CTR vs 목표 Target CTR<br><span style='font-size:12px;color:gray'>마우스를 올려보세요</span>", 'font': {'size': 16}},
                        gauge={'axis': {'range': [None, max(target_ctr * 1.5, current_ctr * 1.2)]}, 'bar': {'color': "green" if current_ctr >= target_ctr else "red"},
                               'threshold': {'line': {'color': BRAIN_CUBE_TEXT, 'width': 4}, 'thickness': 0.75, 'value': target_ctr}}
                    ))
                    fig_ctr.add_trace(go.Pie(
                        values=[1], textinfo='none', hoverinfo='text',
                        hovertext=f"🎯 현재 CTR: {current_ctr:.2f} %<br>🚩 목표 CTR: {target_ctr:.2f} %<br>💡 목표 대비: {current_ctr - target_ctr:.2f} %p ({'초과 달성 ✅' if current_ctr >= target_ctr else '미달 🚨'})",
                        hole=0.5, marker=dict(colors=['rgba(0,0,0,0)']), 
                        domain={'x': [0, 1], 'y': [0, 0.75]}, showlegend=False
                    ))
                    fig_ctr.update_layout(height=300, margin=dict(l=20, r=20, t=50, b=20), hoverlabel=dict(bgcolor="white", font_size=14))
                    g_cols[col_idx].plotly_chart(fig_ctr, use_container_width=True)
                    col_idx += 1
            
            m_col1, m_col2, m_col3, m_col4 = st.columns(4)
            m_col1.metric("총 집행금액", f"{total_spent:,.0f} 원")
            if has_imp: m_col2.metric("총 노출/발송", f"{total_imp:,.0f} 건")
            if has_clicks: m_col3.metric("총 클릭수", f"{total_clicks:,.0f} 회")
            
            metric_label = f"평균 {main_metric}" if main_metric in ['CPA', 'CTR', 'CPC'] else "효율 (데이터 부족)"
            tooltip_msg = "단순 평균이 아닌 '가중 평균(전체 볼륨 대비 성과)' 수치로, 광고주 보고 시 오차가 없는 정석 데이터입니다." if main_metric in ['CTR', 'CPA'] else None
            
            if bench_metric_delta: m_col4.metric(metric_label, efficiency_value, delta=bench_metric_delta, delta_color="inverse" if main_metric == "CPA" else "normal", help=tooltip_msg)
            else: m_col4.metric(metric_label, efficiency_value, help=tooltip_msg)
            
            if col_creative != "(항목 없음)":
                valid_creative_data = df[col_creative].astype(str).replace(['nan', 'NaN', 'None', '', '-', '0'], pd.NA).dropna()
                if valid_creative_data.empty:
                    col_creative = "(항목 없음)"; st.info("💡 선택된 '소재명' 칸이 비어있는 것을 감지하여 자동으로 **[매체별 분석]**으로 전환했습니다.")

            is_creative_mode = col_creative != "(항목 없음)"
            group_target = col_creative if is_creative_mode else col_media
            
            agg_dict = {col_cost: 'sum'}
            if has_imp: agg_dict[col_imp] = 'sum'
            if has_clicks: agg_dict[col_clicks] = 'sum'
            if has_db: agg_dict[col_db] = 'sum'

            ranking_df = filtered_df.groupby(group_target).agg(agg_dict).reset_index()
            if has_db: ranking_df['CPA'] = ranking_df.apply(lambda r: r[col_cost]/r[col_db] if r[col_db]>0 else 0, axis=1)
            if has_clicks: ranking_df['CPC'] = ranking_df.apply(lambda r: r[col_cost]/r[col_clicks] if r[col_clicks]>0 else 0, axis=1)
            if has_imp and has_clicks: ranking_df['CTR'] = ranking_df.apply(lambda r: (r[col_clicks]/r[col_imp])*100 if r[col_imp]>0 else 0, axis=1)

            valid_data = ranking_df[ranking_df[col_db] > 0] if main_metric == 'CPA' else ranking_df[ranking_df[col_clicks] > 0] if main_metric in ['CPC', 'CTR'] else ranking_df
            top_data = valid_data.sort_values(by=main_metric, ascending=is_ascending).head(3)
            summary_text = ranking_df.to_string(index=False)

            st.divider(); st.subheader("📊 다각도 퍼포먼스 시각화")
            tab1, tab2, tab3, tab4 = st.tabs(["🍩 파이/도넛 (비중)", "🫧 버블 맵 (포지셔닝)", "⚖️ 평균 비교 (막대)", "🗓️ 요일별 최적화 (히트맵)"])
            
            with tab1:
                pie_col1, pie_col2 = st.columns(2)
                with pie_col1:
                    if col_cost in ranking_df.columns:
                        fig_pie_cost = px.pie(ranking_df, values=col_cost, names=group_target, hole=0.4, title=f"💰 {group_target}별 집행금액 비중", color_discrete_sequence=CHART_COLORS)
                        fig_pie_cost.update_traces(textposition='inside', textinfo='percent+label'); st.plotly_chart(fig_pie_cost, use_container_width=True); st.session_state.saved_charts['집행금액 비중 (도넛)'] = fig_pie_cost
                with pie_col2:
                    if has_db: perf_col, perf_title = col_db, "DB/전환수"
                    elif has_clicks: perf_col, perf_title = col_clicks, "클릭수"
                    elif has_imp: perf_col, perf_title = col_imp, "발송/노출수"
                    else: perf_col = "(항목 없음)"
                    if perf_col != "(항목 없음)" and perf_col in ranking_df.columns:
                        fig_pie_perf = px.pie(ranking_df, values=perf_col, names=group_target, hole=0.4, title=f"🎯 {group_target}별 {perf_title} 획득 비중", color_discrete_sequence=CHART_COLORS)
                        fig_pie_perf.update_traces(textposition='inside', textinfo='percent+label'); st.plotly_chart(fig_pie_perf, use_container_width=True); st.session_state.saved_charts[f'{perf_title} 획득 비중 (도넛)'] = fig_pie_perf

            with tab2:
                if main_metric in ['CPA', 'CTR', 'CPC'] and len(ranking_df) > 0:
                    y_axis_col = col_db if has_db else (col_clicks if has_clicks else col_imp)
                    y_axis_name = "DB/전환수 (볼륨)" if has_db else ("클릭수 (볼륨)" if has_clicks else "발송/노출수 (볼륨)")
                    if y_axis_col in ranking_df.columns:
                        valid_bubble = ranking_df[ranking_df[y_axis_col] > 0]
                        if not valid_bubble.empty:
                            fig_bubble = px.scatter(valid_bubble, x=main_metric, y=y_axis_col, size=col_cost, color=group_target, hover_name=group_target, size_max=40, title=f"🫧 효율 vs 볼륨 포트폴리오 맵", labels={main_metric: f"효율 ({main_metric})", y_axis_col: y_axis_name}, color_discrete_sequence=CHART_COLORS, template='plotly_white'); st.plotly_chart(fig_bubble, use_container_width=True); st.session_state.saved_charts['효율 vs 볼륨 (버블 차트)'] = fig_bubble

            with tab3:
                fallback_to_campaign_avg = False
                if temp_bench is not None and main_metric in ['CTR', 'CPA']:
                    curr_media_agg = filtered_df.groupby(col_media).agg(agg_dict).reset_index()
                    if has_db: curr_media_agg['CPA'] = curr_media_agg.apply(lambda r: r[col_cost]/r[col_db] if r[col_db]>0 else 0, axis=1)
                    if has_clicks and has_imp: curr_media_agg['CTR'] = curr_media_agg.apply(lambda r: (r[col_clicks]/r[col_imp])*100 if r[col_imp]>0 else 0, axis=1)
                    bench_media_agg = temp_bench.groupby('매체').agg({'발송건':'sum', '클릭수':'sum', '광고비':'sum', 'DB':'sum'}).reset_index()
                    bench_media_agg['Bench_CTR'] = bench_media_agg.apply(lambda r: (r['클릭수']/r['발송건'])*100 if r['발송건']>0 else 0, axis=1)
                    bench_media_agg['Bench_CPA'] = bench_media_agg.apply(lambda r: r['광고비']/r['DB'] if r['DB']>0 else 0, axis=1)
                    comp_df = pd.merge(curr_media_agg, bench_media_agg[['매체', f'Bench_{main_metric}']], left_on=col_media, right_on='매체', how='inner').fillna(0)
                    if not comp_df.empty:
                        melt_df = comp_df.melt(id_vars=[col_media], value_vars=[main_metric, f'Bench_{main_metric}'], var_name='구분', value_name=f'{main_metric} 수치')
                        melt_df['구분'] = melt_df['구분'].replace({main_metric: '이번 캠페인', f'Bench_{main_metric}': '26년 자사 평균'})
                        fig_comp = px.bar(melt_df, x=col_media, y=f'{main_metric} 수치', color='구분', barmode='group', title=f"⚖️ 매체별 {main_metric} 성과 비교 (vs 26년 평균)", color_discrete_sequence=[BRAIN_CUBE_ORANGE, '#CCCCCC'], template='plotly_white'); st.plotly_chart(fig_comp, use_container_width=True); st.session_state.saved_charts['26년 자사 평균 비교 (막대그래프)'] = fig_comp
                    else: fallback_to_campaign_avg = True
                else: fallback_to_campaign_avg = True
                    
                if fallback_to_campaign_avg:
                    if main_metric in ['CTR', 'CPA'] and len(valid_data) > 0:
                        avg_val = current_cpa if main_metric == 'CPA' else current_ctr
                        comp_df2 = valid_data.copy()
                        comp_df2[f'Campaign_Avg_{main_metric}'] = avg_val
                        melt_df2 = comp_df2.melt(id_vars=[group_target], value_vars=[main_metric, f'Campaign_Avg_{main_metric}'], var_name='구분', value_name=f'{main_metric} 수치')
                        melt_df2['구분'] = melt_df2['구분'].replace({main_metric: f'각 {group_target} 실적', f'Campaign_Avg_{main_metric}': '캠페인 전체 평균'})
                        fig_comp2 = px.bar(melt_df2, x=group_target, y=f'{main_metric} 수치', color='구분', barmode='group', title=f"⚖️ {group_target} vs 캠페인 전체 평균 ({main_metric})", color_discrete_sequence=[BRAIN_CUBE_ORANGE, '#8C564B'], template='plotly_white'); fig_comp2.add_hline(y=avg_val, line_dash="dot", line_color="#8C564B", annotation_text="전체 평균"); st.plotly_chart(fig_comp2, use_container_width=True); st.session_state.saved_charts['캠페인 평균 비교 (막대그래프)'] = fig_comp2

            with tab4:
                heat_df = filtered_df.copy()
                heat_df['요일번호'] = heat_df[col_date].dt.dayofweek
                heat_df['요일'] = heat_df['요일번호'].map({0:'월', 1:'화', 2:'수', 3:'목', 4:'금', 5:'토', 6:'일'})
                heat_agg = {col_cost: 'sum'}
                if has_imp: heat_agg[col_imp] = 'sum'
                if has_clicks: heat_agg[col_clicks] = 'sum'
                if has_db: heat_agg[col_db] = 'sum'
                heatmap_data = heat_df.groupby(['요일번호', '요일', col_media]).agg(heat_agg).reset_index()
                if not heatmap_data.empty:
                    if has_db: 
                        heatmap_data['CPA'] = heatmap_data.apply(lambda r: r[col_cost]/r[col_db] if r[col_db]>0 else 0, axis=1)
                        heat_trend, heat_label, color_scale, text_format = 'CPA', "CPA (원)", "Oranges_r", '.0f'
                    elif has_clicks and has_imp: 
                        heatmap_data['CTR'] = heatmap_data.apply(lambda r: (r[col_clicks]/r[col_imp])*100 if r[col_imp]>0 else 0, axis=1)
                        heat_trend, heat_label, color_scale, text_format = 'CTR', "CTR (%)", "Oranges", '.2f'
                    elif has_clicks: 
                        heatmap_data['CPC'] = heatmap_data.apply(lambda r: r[col_cost]/r[col_clicks] if r[col_clicks]>0 else 0, axis=1)
                        heat_trend, heat_label, color_scale, text_format = 'CPC', "CPC (원)", "Oranges_r", '.0f'
                    else: heat_trend, heat_label, color_scale, text_format = col_cost, "집행금액 (원)", "Oranges", '.0f'
                    fig_heat = px.density_heatmap(heatmap_data, x='요일', y=col_media, z=heat_trend, histfunc='sum', text_auto=text_format, title=f"🗓️ 요일 및 매체별 {heat_label} 최적화 히트맵 (진한 색이 효율 좋음)", color_continuous_scale=color_scale, category_orders={"요일": ["월", "화", "수", "목", "금", "토", "일"]}); st.plotly_chart(fig_heat, use_container_width=True)

            st.divider(); st.subheader("📈 일별 퍼포먼스 트렌드")
            daily_data = filtered_df.groupby([col_date, col_media]).agg(agg_dict).reset_index()
            if has_db: daily_data['CPA'] = daily_data.apply(lambda r: r[col_cost]/r[col_db] if r[col_db]>0 else 0, axis=1); trend_col, y_label = 'CPA', "CPA (원)"
            elif has_clicks and has_imp: daily_data['CTR'] = daily_data.apply(lambda r: (r[col_clicks]/r[col_imp])*100 if r[col_imp]>0 else 0, axis=1); trend_col, y_label = 'CTR', "클릭률(CTR) (%)"
            elif has_clicks: trend_col, y_label = col_clicks, "클릭수 (회)"
            elif has_imp: trend_col, y_label = col_imp, "발송/노출수 (건)"
            else: trend_col, y_label = col_cost, "집행금액 (원)"
            if trend_col in daily_data.columns:
                fig_line = px.line(daily_data, x=col_date, y=trend_col, color=col_media, markers=True, labels={col_date: '날짜', trend_col: y_label, col_media: '매체명'}, template='plotly_white', color_discrete_sequence=CHART_COLORS); fig_line.update_traces(marker=dict(size=8)); st.plotly_chart(fig_line, use_container_width=True); st.session_state.saved_charts['일별 퍼포먼스 트렌드 (라인 차트)'] = fig_line

            st.divider(); matched_image_parts_for_gemini = []
            display_metric_name = main_metric if main_metric in ['CPA', 'CTR', 'CPC'] else "집행금액"
            if is_creative_mode: st.subheader(f"🔥 Top 성과 소재 분석 ({display_metric_name} 기준)")
            else: st.subheader(f"🏆 Top 성과 매체 분석 ({display_metric_name} 기준)")
            
            with st.expander(f"📊 {group_target} 전체 랭킹 데이터 펼쳐보기"): st.dataframe(ranking_df.sort_values(by=main_metric, ascending=is_ascending), use_container_width=True)

            if len(top_data) > 0:
                cols = st.columns(len(top_data))
                for i, (index, row) in enumerate(top_data.iterrows()):
                    item_name = str(row[group_target])
                    with cols[i]:
                        st.markdown(f"<h3 style='color: {BRAIN_CUBE_ORANGE};'>Top {i+1}</h3>", unsafe_allow_html=True)
                        if is_creative_mode and item_name in image_files:
                            img_file = image_files[item_name]; st.image(img_file, use_container_width=True); matched_image_parts_for_gemini.append({"mime_type": "image/jpeg" if item_name.endswith(('.jpg', '.jpeg')) else "image/png", "data": img_file.getvalue()})
                        else: st.info(f"✨ **{item_name}**") 
                        st.write(f"- 집행금액: {row[col_cost]:,.0f}원")
                        if has_imp and col_imp in row: st.write(f"- 노출/발송: {row[col_imp]:,.0f}건")
                        if has_clicks and col_clicks in row: st.write(f"- 클릭수: {row[col_clicks]:,.0f}회")
                        if has_db and col_db in row: st.write(f"- DB/전환: {row[col_db]:,.0f}건")
                        if has_db: st.write(f"- **CPA: <span style='color: {BRAIN_CUBE_ORANGE}; font-weight: bold;'>{row['CPA']:,.0f}원</span>**", unsafe_allow_html=True)
                        if has_imp and has_clicks: st.write(f"- **CTR: <span style='color: {BRAIN_CUBE_ORANGE}; font-weight: bold;'>{row['CTR']:.2f}%</span>**", unsafe_allow_html=True)

            st.divider(); st.subheader("🤖 브레인큐브 AI 종합 퍼포먼스 분석")
            if st.button("브레인큐브 AI 리포트 생성하기"):
                with st.spinner("다각도 시각화 데이터와 26년 평균을 바탕으로 꼼꼼히 분석 중입니다..."):
                    if is_creative_mode: visual_instruction = "첨부된 이미지를 분석하거나 문맥을 파악하여 성과 이유를 설명하세요." if len(matched_image_parts_for_gemini) > 0 else "위 Top 성과를 낸 '텍스트'의 문맥적 매력도를 분석하세요."; target_name = "소재별"
                    else: visual_instruction = "위 Top 매체들의 성과를 분석하고 왜 효율이 좋았는지 요일별 데이터와 타겟팅 관점에서 분석하세요."; target_name = "매체별"
                    prompt_text = f"당신은 '브레인큐브(Brain Cube)'의 최고 퍼포먼스 마케터입니다. 이번 캠페인의 효율은 '{efficiency_value}'로 측정되었으며, 26년 브레인큐브 전체 평균({selected_industry} 기준) 대비 {'좋은' if (main_metric=='CTR' and bench_metric_delta and '+' in bench_metric_delta) or (main_metric=='CPA' and bench_metric_delta and '-' in bench_metric_delta) else '아쉬운'} 성과를 보였습니다.\n[작성 지침] 1. 파일명('{file_name}')에서 광고주를 추출해 전문적인 보고서 제목 작성. 2. 인사말 없이 바로 브레인큐브의 공식 보고서답게 확신에 찬 어조로 작성. 3. {visual_instruction} 4. 26년 평균 대비 잘된 점이나 보완할 점을 언급하며, 향후 최적화 전략 제안.\n[{target_name} 세부 퍼포먼스 데이터]\n{summary_text}"
                    response = model.generate_content([prompt_text] + matched_image_parts_for_gemini); st.session_state.ai_report_content = response.text

            if st.session_state.ai_report_content:
                st.info(st.session_state.ai_report_content)
                def create_word(text, figs):
                    try:
                        from docx import Document; from docx.shared import Inches; from docx.oxml.ns import qn
                        doc = Document(); style = doc.styles['Normal']; style.font.name = '맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        doc.add_heading('Brain Cube AI Performance Report', 0)
                        clean_text = re.sub(r'\*\*', '', text); clean_text = re.sub(r'#+\s*', '', clean_text)
                        for paragraph in clean_text.split('\n'):
                            if paragraph.strip(): doc.add_paragraph(paragraph)
                        if figs:
                            doc.add_page_break(); doc.add_heading('📈 첨부 데이터 시각화 차트', level=1)
                            for title, fig in figs.items():
                                try: doc.add_heading(title, level=2); img_stream = io.BytesIO(fig.to_image(format="png", engine="kaleido")); doc.add_picture(img_stream, width=Inches(6.0))
                                except Exception as e: doc.add_paragraph(f"[{title} 차트 이미지 변환 실패 - {e}]")
                        target_stream = io.BytesIO(); doc.save(target_stream); return target_stream.getvalue()
                    except Exception as e: st.error(f"Word 생성 오류: {e}"); return b""
                st.download_button(label="📝 그래프 포함 리포트 다운로드 (.docx)", data=create_word(st.session_state.ai_report_content, st.session_state.saved_charts), file_name=f"BrainCube_AI_Report_{file_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        else: st.warning("달력에서 시작일과 종료일을 모두 선택해주세요.")
    except Exception as e: st.error(f"분석 중 오류가 발생했습니다: {e}\n\n화면 상단의 [⚙️ 데이터 항목 매칭]이 올바르게 선택되었는지 확인해 주세요!")
else: st.info("캠페인 결과 엑셀 파일을 업로드해주세요.")